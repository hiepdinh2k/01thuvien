
from flask import Flask, render_template, request
import pandas as pd
from docx import Document
import io,csv
from model import *
import os
app = Flask(__name__)

# ==========================Function=================================
def read_docx_tables(filename,tab_id = None, **kwargs):
    def read_docx_tab(tab,**kwargs):
        vf = io.StringIO()
        writer = csv.writer(vf)
        for row in tab.rows:
            writer.writerow(cell.text for cell in row.cells)
        vf.seek(0)
        df = pd.read_csv(vf,**kwargs)
        # df = df.replace(r'\n',' ', regex=True)
        return df
    doc = Document(filename)
    if tab_id is None:
        return[read_docx_tab(tab,**kwargs) for tab in doc.tables]
    else:
        try:
            return read_docx_tab(doc.tables[tab_id],**kwargs)
        except IndexError:
            print("error: specified [tab_id]: {} does not exist.".format(tab_id))




#===========================Flask==================================
@app.route('/', methods=['GET','POST'])
def upload():
    if request.method == 'POST':
        file = request.files['file1']
        upload_dir = f"data/ProjectName/{request.environ['REMOTE_ADDR']}"
        os.makedirs(upload_dir, exist_ok=True)
        objects = os.listdir(upload_dir)
        files_file = [f for f in objects if os.path.isfile(os.path.join(upload_dir, f))]
        for f in files_file:
            os.remove(os.path.join(upload_dir, f))
        file.save(os.path.join(upload_dir, file.filename))
        df = read_docx_tables(file)
        tb3 = chuyen_tb3(df[2])
        tb2 = table2(df[1])
        tb1 = table1(df[0])
        
        return render_template('index3.html', tables=
        [tb1.to_html(classes='table table-hover' , index=False , table_id='tb1-resuly'),
         tb2.to_html(classes='table table-hover ' , index=False , table_id='tb2-resuly'),
         tb3.to_html(classes='table table-hover' , header=False , table_id='tb3-resuly')])  
        
    return render_template('index3.html')
 

if __name__=='__main__':
    app.run(debug=True ,port= 2345,host='0.0.0.0')